import fitz  # PyMuPDF
from PIL import Image, ImageEnhance, ImageFilter
from docx import Document
from docx.shared import Inches
from io import BytesIO

def pdf_to_word(pdf_path, output_docx, dpi=600, contrast_factor=1.5, sharpness_factor=2.0, clarity_factor=1.5):
    # Open the PDF file
    pdf_document = fitz.open(pdf_path)

    # Load the first page as an image with the specified DPI and apply enhancements
    page = pdf_document.load_page(0)
    image = page.get_pixmap(matrix=fitz.Matrix(dpi / 72, dpi / 72))
    
    # Convert to a Pillow image
    image = Image.frombytes("RGB", [image.width, image.height], image.samples)

    # Enhance contrast
    enhancer = ImageEnhance.Contrast(image)
    image = enhancer.enhance(contrast_factor)

    # Enhance sharpness
    image = image.filter(ImageFilter.SHARPEN)

    # Enhance clarity
    enhancer = ImageEnhance.Sharpness(image)
    image = enhancer.enhance(clarity_factor)

    # Create a Word document
    doc = Document()

    # Add a title to the document
    doc.add_heading('Scanned Image', 0)

    # Add the scanned image to the document
    img_buffer = BytesIO()
    image.save(img_buffer, format="PNG")
    doc.add_picture(img_buffer, width=Inches(6.0))

    # Save the Word document
    doc.save(output_docx)

    # Close the PDF file
    pdf_document.close()

if __name__ == "__main__":
    pdf_file = "sample_pdf\\0.pdf"
    output_docx_file = "output_document.docx"

    # Convert PDF to Word document directly
    pdf_to_word(pdf_file, output_docx_file, dpi=1200, contrast_factor=1.5, sharpness_factor=2.0, clarity_factor=1.5)
